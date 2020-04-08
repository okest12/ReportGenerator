import re
import docx
import os
import win32com
from win32com.client import Dispatch

def print_by_docx(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            print(run.text)

def replace_by_docx():
    doc = docx.Document(r"C:\01-code\2-2019年企业所得税审核报告及说明范本.docx")
    print_by_docx(doc)
    for para in doc.paragraphs:
        m = re.search('S[0-9]+[A-Z][0-9]+', para.text)
        while m:
            para.text = para.text.replace(m.group(0), "123456")
            m = re.search('S[0-9]+[A-Z][0-9]+', para.text)
    doc.save(r"C:\01-code\2-2019年企业所得税审核报告及说明范本_new1.docx")

def replace_by_docx_run():
    doc = docx.Document(r"C:\01-code\2-2019年企业所得税审核报告及说明范本.docx")
    print_by_docx(doc)
    for para in doc.paragraphs:
        for run in para.runs:
            m = re.search('S[0-9]+[A-Z][0-9]+', run.text)
            if m:
                run.text = "12345"
    doc.save(r"C:\01-code\2-2019年企业所得税审核报告及说明范本_new1.docx")


class RemoteWord:
    def __init__(self, filename=None):

        self.xlApp = win32com.client.Dispatch('Word.Application')

        self.xlApp.Visible = 0  # 后台运行，不显示
        self.xlApp.DisplayAlerts = 0 # 不警告
        if filename:
            self.filename = filename
            if os.path.exists(self.filename):
                self.doc = self.xlApp.Documents.Open(filename)
            else:
                self.doc = self.xlApp.Documents.Add() # 创建新的文档
                self.doc.SaveAs(filename)
        else:
            self.doc = self.xlApp.Documents.Add()
            self.filename = ''

    def add_doc_end(self, string):
        rangee = self.doc.Range()
        rangee.InsertAfter('\n' + string)


    def add_doc_start(self, string):
        rangee = self.doc.Range(0, 0)
        rangee.InsertBefore(string + '\n')

    def insert_doc(self, insertPos, string):
        rangee = self.doc.Range(0, insertPos)
        if insertPos == 0:
            rangee.InsertAfter(string)
        else:
            rangee.InsertAfter('\n' + string)

    def replace_doc(self, string, new_string):
        self.xlApp.Selection.Find.ClearFormatting()
        self.xlApp.Selection.Find.Replacement.ClearFormatting()
        # (string--搜索文本,
        # True--区分大小写,
        # True--完全匹配的单词，并非单词中的部分（全字匹配）,
        # True--使用通配符,
        # True--同音,
        # True--查找单词的各种形式,
        # True--向文档尾部搜索,
        # 1,
        # True--带格式的文本,
        # new_string--替换文本,
        # 2--替换个数（全部替换）
        self.xlApp.Selection.Find.Execute(string, False, False, False, False, False, True, 1, True, new_string, 2)

    def replace_docs_re(self, string, new_string):
        '''采用通配符匹配替换'''

        self.xlApp.Selection.Find.ClearFormatting()
        self.xlApp.Selection.Find.Replacement.ClearFormatting()
        self.xlApp.Selection.Find.Execute(string, False, False, True, False, False, False, 1, False, new_string, 2)

    def save(self):
        self.doc.Save()


    def save_as(self, filename):
        self.doc.SaveAs(filename)


    def close(self):
        #self.save()
        self.doc.SaveAs(r"C:\01-code\2-2019年企业所得税审核报告及说明范本_new.docx")
        self.xlApp.Documents.Close()
        self.xlApp.Quit()

def replace_by_win32_com():
    path = r"C:\01-code\2-2019年企业所得税审核报告及说明范本.docx"
    doc = RemoteWord(path)  # 初始化一个doc对象
    # 这里演示替换内容，其他功能自己按照上面类的功能按需使用

    doc.replace_doc('S5C22', '12345')  # 替换文本内容
    doc.replace_doc('S4D32', '24680')  # 替换文本内容
    #doc.replace_doc('．', '.')  # 替换．为.
    #doc.replace_doc('\n', '')  # 去除空行
    #doc.replace_doc('o', '0')  # 替换o为0
    # doc.replace_docs('([0-9])@[、,，]([0-9])@', '\1.\2')  使用@不能识别改用{1,}，\需要使用反斜杠转义
    #doc.replace_docs('([0-9]){1,}[、,，．]([0-9]){1,}', '\\1.\\2')  # 将数字中间的，,、．替换成.
    #doc.replace_docs('([0-9]){1,}[旧]([0-9]){1,}', '\\101\\2')  # 将数字中间的“旧”替换成“01”
    doc.close()

replace_by_win32_com()